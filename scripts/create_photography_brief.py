from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUT = Path('deliverables/ChristianScienceAurora_Phase1_Photography_and_AI_Image_Prompts.docx')

TEAL = '167C80'
DEEP_TEAL = '0D5559'
PERI = '7379B8'
DEEP_PERI = '53598E'
MIST = 'DDEEEF'
PERI_MIST = 'E8E9F5'
WARM = 'FCFBF7'
STONE = 'F1EFE9'
CHARCOAL = '263437'
GOLD = 'D6A64B'
WHITE = 'FFFFFF'

assets = [
    dict(group='Homepage', id='campaign-public-talk', placement='Home — administrator-controlled campaign banner', priority='Conditional / campaign-driven', ratio='Desktop 16:9; mobile 4:5 or 3:4 crop', label='Featured event or campaign image', objective='Create a strong, timely visual that gives a confirmed public talk, lecture, family program, or community event immediate prominence without overpowering the welcome message.', brief='Photograph the actual event speaker, welcome table, or event environment. Favor a candid side angle, natural room light, and visible but not dominant audience context. Leave clean negative space for a headline and date. Capture separate horizontal and vertical compositions rather than relying on one crop.', prompt='Using the supplied approved event and location reference photographs, create an ultra-realistic editorial event photograph for Christian Science Aurora. Preserve the real room, architecture, lighting, speaker appearance, and event details exactly. Show a welcoming public talk in progress from a respectful side angle, with a small attentive audience softly present in the background. Natural window and practical light, realistic skin texture, calm documentary mood, warm white and subtle teal/periwinkle visual harmony, generous negative space on the left for website headline copy, no staged smiles, no dramatic spotlight, no artificial glow. 16:9 horizontal composition, high resolution, believable local church setting.', negative='No invented speaker, false event details, decorative cross, halo, glowing hands, dramatic clouds, corporate conference staging, oversaturated teal, readable copyrighted pages, third-party logos, or identifiable minors.', alt='A speaker presenting to an audience at a Christian Science Aurora public event.', release='Speaker/event-media permission, adult participant releases where recognizable, venue approval, and no youth faces without guardian authorization.'),
    dict(group='Homepage', id='home-welcome', placement='Home — primary welcome hero', priority='Launch critical', ratio='Desktop 4:3; mobile 4:5 focal-safe crop', label='Local church welcome', objective='Establish immediately that the church is peaceful, modern, local, and open to first-time visitors.', brief='Use the actual entrance, foyer, or sanctuary in soft morning light. One or two naturally moving adults may be included if released, but architecture and a clear sense of arrival should remain primary. Keep the frame uncluttered and leave copy-safe space.', prompt='Use supplied reference photos of First Church of Christ, Scientist, Aurora as strict architectural references. Create an ultra-realistic, high-end editorial photograph of the real church entrance in soft Colorado morning light. Preserve exact building proportions, doors, windows, materials, signage, landscaping, and accessibility features; do not add or remove architecture. The entrance is open and warmly lit, with one adult visitor approaching naturally from a three-quarter rear angle. Calm, truthful documentary photography, realistic contrast, subtle depth of field, warm white highlights, restrained teal and periwinkle color harmony, ample negative space for website copy, 4:3 horizontal composition with a mobile-safe center subject.', negative='No invented architecture, dramatic sky replacement, sun rays, artificial halo, glowing doorway, staged corporate smile, crowds, decorative religious symbols, heavy HDR, text overlays, or false signage.', alt='Welcoming view of the entrance to First Church of Christ, Scientist, Aurora.', release='Property approval; model release for any recognizable person; verify signage and accessible-route accuracy.'),
    dict(group='Homepage', id='new-here', placement='Home — New Here / visitor reassurance section', priority='Launch critical', ratio='4:3', label='Natural arrival and program handoff', objective='Reduce first-visit anxiety by showing a simple, respectful welcome rather than a staged greeting.', brief='Photograph one adult greeter offering a service program to one adult visitor near the real entrance or foyer. Capture an in-between moment with natural body language. Keep faces secondary or photograph from a respectful side angle.', prompt='Using supplied reference photographs of the actual Aurora church foyer and approved adult models, create an ultra-realistic documentary image of a simple first-time welcome. An adult greeter naturally offers a service program to an adult visitor near the entrance; both are relaxed, respectful, and not posing for the camera. Soft natural daylight, truthful skin texture, understated clothing, uncluttered background, modern editorial composition, realistic Colorado church setting, subtle teal/periwinkle accents in the environment, 4:3 crop with open space around the interaction.', negative='No exaggerated handshake, staged grin, name tags unless real, sales-like interaction, crowds, glowing effects, praying pose, decorative cross, readable protected publication pages, or identifiable minors.', alt='A visitor receiving a service program near the church entrance.', release='Adult model releases and confirmation that any visible program cover or text may be shown.'),
    dict(group='Homepage', id='latest-sermon', placement='Home — latest sermon feature', priority='Launch support', ratio='16:9', label='Sanctuary and sermon media', objective='Provide a calm, permission-aware visual for approved sermon recordings or summaries.', brief='Photograph the sanctuary or Readers’ platform before a service. Use a wide or medium detail with no readable protected text. The setting should feel ready, ordered, and quiet.', prompt='Using reference photographs of the actual Christian Science Aurora sanctuary, create an ultra-realistic editorial photograph of the Readers’ platform prepared before a service. Preserve exact architecture and furnishings. Show approved books closed or open with all page text unreadable, balanced natural daylight, realistic wood and fabric textures, clean composition, quiet expectancy, no people, no artificial glow, 16:9 frame with room for a media play control and title overlay.', negative='No readable Bible Lesson text, altered architecture, halo, cross added as decoration, cinematic spotlight, floating dust beams, heavy vignette, publication logos enlarged, or fake audience.', alt='The sanctuary and Readers’ platform prepared for a Christian Science service.', release='Confirm visual use of books, platform details, trademarks, and any publication covers.'),
    dict(group='Homepage', id='home-reading-room', placement='Home — Reading Room feature card', priority='Launch critical', ratio='4:3', label='Reading Room shelves and quiet table', objective='Make the Reading Room feel accessible, useful, and calm rather than like a retail store.', brief='Photograph the actual Reading Room with shelves, a clear reading table, and natural light. Keep authorized covers small and contextual. Include a staff member helping a visitor only if releases are complete.', prompt='Use supplied photographs of the real Christian Science Reading Room as strict references. Create an ultra-realistic editorial interior photograph showing a quiet reading table, organized shelves, and soft natural daylight. Preserve exact room layout, furniture, materials, and approved signage. The scene feels open, uncluttered, and welcoming, with subtle human warmth but no posed subjects. Publication covers are present only as small contextual objects and are not readable. 4:3 horizontal composition, realistic color, warm white surfaces, restrained teal/periwinkle accents.', negative='No bookstore merchandising wall, fake publication covers, enlarged logos, staged salesperson, artificial glow, overexposed windows, dramatic shadows, or invented room features.', alt='Reading Room shelves and a quiet table in natural light.', release='Location approval and permission review for every recognizable cover, logo, sign, or person.'),
    dict(group='Homepage', id='home-sunday-school', placement='Home — Sunday School feature card', priority='Launch critical', ratio='4:3', label='Hands-on children’s activity', objective='Show learning, creativity, and curiosity while protecting children’s identities.', brief='Photograph hands, materials, and teacher guidance rather than faces. Use original worksheets or blank activity materials. Keep the scene bright, natural, and genuinely age-appropriate.', prompt='Create an ultra-realistic, respectful editorial photograph of a Sunday School table activity, using supplied room references and approved participants. Frame only children’s hands and forearms as they color an original gratitude activity beside a teacher’s guiding hand. No faces, names, school logos, or identifying details. Natural window light, realistic paper and crayon texture, warm white tabletop, subtle teal, periwinkle, and restrained gold materials, joyful but calm composition, 4:3 crop, high resolution.', negative='No identifiable child faces, surnames, name tags, school uniforms, location metadata, generic praying hands, adult-looking child hands, fake Bible text, branded worksheets, glittery fantasy effects, or separate Sunday School logo.', alt='Children’s hands working on a Sunday School activity with a teacher nearby.', release='Guardian releases if any child is recognizable; metadata stripping; youth editorial and compliance approval.'),
    dict(group='Homepage', id='home-article', placement='Home — featured article card', priority='Launch support', ratio='4:3', label='Quiet Aurora reflection moment', objective='Support editorial content with a grounded local image that avoids wellness clichés.', brief='Photograph an adult reading near a window, a quiet architectural detail, or a recognizable Aurora landscape detail. Keep the moment observational, not devotional or posed.', prompt='Create an ultra-realistic editorial photograph of a quiet reading and reflection moment in Aurora, Colorado. An adult sits near a naturally lit window reading a plain, unbranded book, photographed from the side with a calm, thoughtful posture. Realistic home or Reading Room environment, Colorado daylight, subtle local landscape visible out of focus, warm white palette with restrained teal and periwinkle accents, no staged smile, no spiritual special effects, 4:3 composition with copy-safe space.', negative='No meditation pose, yoga styling, glowing hands, halo, dramatic mountains unless geographically accurate, fake religious text, visible brand logos, medical imagery, or guaranteed-healing symbolism.', alt='A quiet reading and reflection moment in natural light.', release='Adult model and location release; no readable protected content.'),
    dict(group='Homepage', id='local-community', placement='Home — local community story', priority='Launch critical', ratio='16:9', label='Wide church and Aurora context', objective='Place the church clearly in Aurora and communicate long-term local presence.', brief='Use a wide exterior showing the actual building, landscaping, sidewalk, and everyday neighborhood context. Include natural human activity only as secondary, released subjects.', prompt='Using supplied reference photographs, create an ultra-realistic wide editorial photograph of First Church of Christ, Scientist, Aurora in its real neighborhood setting. Preserve exact architecture, signage, landscaping, parking, and surrounding context. Clear natural Colorado daylight, a few distant adults entering or leaving naturally, realistic vehicles without readable plates, calm community atmosphere, no dramatic sky, no architectural invention, 16:9 panoramic composition with balanced negative space.', negative='No altered building, added steeple, decorative cross, dramatic sunset, storm clouds, crowds, rally atmosphere, fake sign, visible license plates, or artificial lens flare.', alt='First Church of Christ, Scientist, Aurora in its local community setting.', release='Property approval, releases for recognizable people, blur or avoid plates, and verify signage.'),

    dict(group='Visit and Services', id='visit-entrance', placement='Plan Your Visit — page hero', priority='Launch critical', ratio='16:9 or 4:3', label='Church entrance and arrival', objective='Show exactly where a first-time visitor should go and make the approach feel simple.', brief='Photograph the real approach from a visitor’s eye level. Include the walkway, door, sign, and accessibility features. Avoid decorative staging.', prompt='Using supplied reference photographs of the actual Aurora church, create an ultra-realistic visitor-eye-level photograph of the main entrance and walkway in soft morning light. Preserve exact doors, windows, sign, paving, railings, ramps, and landscaping. The path to the entrance is visually clear and unobstructed. No people required. Natural Colorado color and contrast, warm and peaceful but fully documentary, 16:9 composition with a 4:3 safe crop.', negative='No invented ramp or entrance, altered sign, dramatic sunburst, open door glow, decorative religious symbol, flowers added beyond reality, fake parking signs, or heavy HDR.', alt='Main entrance and walkway at First Church of Christ, Scientist, Aurora.', release='Verify the photographed entrance is the preferred public entrance and accessibility details are accurate.'),
    dict(group='Visit and Services', id='directions-exterior', placement='Directions and Parking — wide orientation image', priority='Launch critical', ratio='16:9', label='Wide exterior and parking approach', objective='Help visitors recognize the property from the street and understand the driveway approach.', brief='Shoot from a legal, safe location showing street approach, driveway, building, and parking relationship. Capture multiple traffic-direction angles if useful.', prompt='Use supplied real-location reference images to create an ultra-realistic wide street-approach photograph of First Church of Christ, Scientist, Aurora. Preserve the exact road relationship, driveway, parking entry, building architecture, landscaping, and signage. Clear midday or late-morning visibility, realistic Colorado weather, no dramatic clouds, no traffic manipulation, 16:9 landscape composition with the driveway visually easy to identify.', negative='No invented turn lane, fake road sign, altered driveway, empty impossible roadway, dramatic sky replacement, readable plates, added steeple, or misleading parking layout.', alt='Wide view of the Aurora church, driveway, parking approach, and main entrance.', release='Confirm safe/legal capture position; obscure license plates; verify directions before publication.'),
    dict(group='Visit and Services', id='parking-diagram', placement='Directions and Parking — annotated guide', priority='Launch critical after church confirmation', ratio='4:3', label='Parking and entrance guide', objective='Provide a factual, accessible orientation graphic rather than a decorative image.', brief='Create from an actual aerial/property plan or accurately stitched elevated photograph. Mark visitor entrance, accessible spaces, preferred route, and overflow only after verification. Use clear text and high contrast.', prompt='Create a clean, accessible 4:3 site-orientation diagram based strictly on the supplied verified property map and photographs for First Church of Christ, Scientist, Aurora. Use a warm white background, simplified charcoal building footprint, Aurora teal route line, periwinkle parking zones, and restrained Morning Gold only for the visitor entrance marker. Label Main Entrance, Accessible Parking, Visitor Parking, and Preferred Walking Route only where confirmed. Large readable Inter-style labels, no decorative icons, no invented spaces, north arrow and street name included.', negative='No fictional parking spaces, inaccurate accessible route, tiny text, color-only meaning, third-party map branding, decorative cross, gradients, or personal vehicle details.', alt='Parking areas, accessible route, and visitor entrance at the Aurora church.', release='Church facilities and accessibility review; verify every label and route.'),
    dict(group='Visit and Services', id='greeter-program', placement='What to Expect — Arrive / Connect section', priority='Launch support', ratio='4:3', label='Natural visitor welcome', objective='Demonstrate that newcomers can enter quietly and receive practical help without pressure.', brief='Capture an authentic program handoff or brief conversation near the entrance. Avoid handshakes or direct-to-camera posing.', prompt='Using the real church entrance or foyer references and released adult participants, create an ultra-realistic candid photograph of an adult visitor being offered a service program by an adult greeter. Side-angle documentary framing, natural body language, soft morning daylight, uncluttered environment, simple respectful clothing, no direct eye contact with camera, 4:3 composition, truthful textures and skin tones.', negative='No staged handshake, exaggerated smile, sales interaction, name tag unless real, glowing light, praying pose, child subject, readable personal information, or altered architecture.', alt='A visitor receiving a service program near the church entrance.', release='Adult releases; verify program imagery/text permissions.'),
    dict(group='Visit and Services', id='sanctuary-wide', placement='Services — landing-page hero', priority='Launch critical', ratio='16:9', label='Sanctuary from a visitor’s perspective', objective='Let newcomers see the room before arriving and understand that the environment is calm and approachable.', brief='Shoot from the rear or center aisle at seated eye level. Show room proportions, seating, windows, platform, and music detail accurately.', prompt='Using supplied reference photographs of the actual sanctuary, create an ultra-realistic wide interior photograph from a visitor’s seated or center-aisle perspective. Preserve exact seating, platform, windows, organ or music area, walls, materials, lighting, and proportions. Balanced natural daylight, clean but not artificially perfect, welcoming and quiet, no people, no added symbols, 16:9 composition with straight architectural lines and realistic dynamic range.', negative='No architectural redesign, added cross, halo, glowing platform, theatrical stage lighting, distorted wide-angle walls, fake flowers, readable protected pages, or oversized trademark.', alt='Interior of the Aurora church sanctuary viewed from the visitor seating area.', release='Property approval and review of any visible marks, books, hymnals, or artwork.'),
    dict(group='Visit and Services', id='sunday-sanctuary-detail', placement='Sunday Service — page hero', priority='Launch support', ratio='16:9', label='Sunday worship environment', objective='Show a prepared worship setting with detail and warmth while remaining factual.', brief='Photograph before attendees arrive. Include orderly seating, hymnals, platform, and one architectural detail. Avoid making protected content readable.', prompt='Create an ultra-realistic editorial photograph of the actual Christian Science Aurora sanctuary prepared for Sunday worship, based strictly on supplied reference images. Warm natural daylight, orderly seating and hymnals, subtle focus toward the Readers’ platform, realistic materials and shadows, no people, no staged religious effects, exact architecture preserved, 16:9 horizontal frame with clean title space.', negative='No readable Bible Lesson pages, invented decor, decorative cross, dramatic spotlight, candles unless actually present, artificial halo, oversaturated colors, or perfect showroom rendering.', alt='Quiet sanctuary prepared for Sunday worship.', release='Review visible hymnals, books, marks, and platform details for permission.'),
    dict(group='Visit and Services', id='wednesday-evening', placement='Wednesday Testimony Meeting — page hero', priority='Launch support', ratio='16:9', label='Church entrance at early evening', objective='Make an evening visit feel calm, safe, and easy to recognize.', brief='Photograph during real blue hour with the actual exterior lights. Keep exposure realistic so the building and walkway remain visible.', prompt='Using actual exterior reference photographs, create an ultra-realistic blue-hour photograph of First Church of Christ, Scientist, Aurora before the Wednesday testimony meeting. Preserve exact architecture, entrance lights, walkway, signage, and landscaping. Realistic early-evening Colorado sky, balanced ambient and practical lighting, calm and safe arrival feeling, no exaggerated glow, no people required, 16:9 composition.', negative='No dramatic sunset, star-filled fantasy sky, glowing doorway, added lights, wet pavement unless real, artificial fog, altered sign, decorative cross, or cinematic lens flare.', alt='Church entrance in the early evening before the Wednesday testimony meeting.', release='Verify lighting and entrance shown are accurate and normally available.'),
    dict(group='Visit and Services', id='sermon-reader-view', placement='Sermon Library — hero and media thumbnails', priority='Launch support', ratio='16:9', label='Readers’ platform and approved books', objective='Create a dignified media image that can be reused across approved sermon records without exposing protected text.', brief='Photograph platform or book detail with pages unreadable. Capture a clean, consistent series at several focal lengths.', prompt='Using reference photographs of the actual Readers’ platform, create an ultra-realistic editorial detail of the platform and approved books before a service. Exact architecture and furnishings preserved, books positioned naturally with page text completely unreadable, soft directional daylight, realistic wood, paper, and fabric texture, quiet and dignified mood, 16:9 media-thumbnail composition with safe center crop.', negative='No legible copyrighted text, enlarged Cross and Crown, added decorative cross, glowing book, sun rays, floating particles, theatrical spotlight, or invented platform elements.', alt='Readers’ platform prepared for a Christian Science service.', release='Permissions review for books, covers, marks, and any visible text.'),

    dict(group='Sunday School', id='sunday-school-hands', placement='Sunday School — primary hero', priority='Launch critical', ratio='16:9', label='Children creating at a table', objective='Communicate curiosity and active learning while minimizing youth-identification risk.', brief='Frame hands, original materials, and teacher guidance. Use age-appropriate diversity without tokenistic staging. Avoid face-forward compositions.', prompt='Create an ultra-realistic editorial photograph in the supplied Sunday School classroom. Show only children’s hands and forearms around a table as they draw, arrange original story cards, and ask questions with a teacher’s hand gently pointing to an activity. No faces or identifying details. Natural daylight, realistic materials, warm white background, restrained teal, periwinkle, and Morning Gold supplies, joyful and thoughtful rather than chaotic, 16:9 composition with large touch-friendly website crop.', negative='No child faces, names, school logos, uniforms, public comments, fake scripture text, branded worksheet, glowing hands, cartoon overlay, separate Sunday School logo, or adult hands posing as children.', alt='Children’s hands working on a Sunday School activity with a teacher nearby.', release='Guardian authorization where recognition is possible; strip metadata; youth compliance review.'),
    dict(group='Sunday School', id='family-arrival', placement='Sunday School Parent Information — hero', priority='Launch support', ratio='16:9', label='Family arriving for Sunday School', objective='Help parents visualize a comfortable first arrival without exposing children’s identities.', brief='Photograph from behind or a non-identifying angle as a parent and child approach the actual entrance. Keep route and door recognizable.', prompt='Using real Aurora church entrance references and released participants, create an ultra-realistic candid photograph of a parent and elementary-age child walking toward the church entrance on a clear Sunday morning. Photograph from behind at a respectful distance so faces are not visible. Preserve exact walkway, door, signage, and landscaping. Natural relaxed movement, simple clothing, warm realistic daylight, 16:9 composition with clear sense of arrival.', negative='No visible child face, surname, school logo, posed handholding, exaggerated smile, glowing entrance, invented building, dramatic sky, or crowd.', alt='A family walking toward the church entrance for Sunday School.', release='Guardian and adult releases; no identifying clothing or metadata; confirm entrance route.'),
    dict(group='Sunday School', id='story-illustration', placement='Children’s Stories — featured original illustration', priority='Launch support', ratio='4:3', label='Original children’s story illustration', objective='Provide a distinctive child-friendly visual within the same brand system without creating a separate logo or mascot.', brief='Commission original editorial illustration. Use simple geometry, human warmth, and clear storytelling. Ensure representation feels natural and respectful.', prompt='Create an original contemporary editorial illustration for Christian Science Aurora children’s stories. A small group of elementary-age children sits with an adult teacher around an open, non-readable storybook, discussing an idea and pointing to simple picture cards. Warm white background, clean geometric shapes, soft hand-drawn texture, Aurora teal and periwinkle as primary accents, Morning Gold used sparingly, diverse children represented naturally, generous whitespace, no logo, no mascot, no religious clip art, 4:3 composition, sophisticated modern children’s publishing quality.', negative='No separate Sunday School logo, cartoon church mascot, decorative cross, halo, glowing hands, stock-vector look, childish rainbow overload, readable copyrighted text, photoreal faces, or brand logos.', alt='Original illustration of children reading and discussing a Bible story together.', release='Original commissioned work with written transfer/license terms and source-file retention.'),

    dict(group='Reading Room and Events', id='reading-room-interior', placement='Reading Room — primary hero', priority='Launch critical', ratio='16:9 or 4:3', label='Reading Room interior', objective='Show the Reading Room as a calm public resource for reading, study, questions, and help.', brief='Photograph wide enough to establish the room. Keep shelves orderly but lived-in. Include staff/visitor interaction only with releases.', prompt='Use supplied photographs of the actual Christian Science Reading Room as strict references. Create an ultra-realistic wide editorial interior showing organized shelves, a quiet reading table, a comfortable chair, and natural daylight. Preserve exact room dimensions, furniture, materials, signage, and layout. Calm public-resource atmosphere rather than retail merchandising. Publication covers remain small and unreadable unless specifically approved. 16:9 composition with a 4:3 safe crop.', negative='No invented room, fake covers, large product display, sales counter emphasis, glowing books, dramatic sun rays, altered signage, people without releases, or excessive visual clutter.', alt='Wide interior view of the Christian Science Reading Room with shelves and a quiet reading table.', release='Location confirmation; cover, logo, sign, and participant permissions.'),
    dict(group='Reading Room and Events', id='product-cover-placeholder', placement='Reading Room Shop — product image system', priority='Required before commerce activation', ratio='2:3', label='Authorized product cover or neutral product image', objective='Create a consistent product presentation without misusing copyrighted covers or official marks.', brief='Use exact publisher-supplied cover files when authorization permits. Never reconstruct or alter a cover with AI. When permission is pending, use a neutral typographic product card with title, author, format, and “image pending authorization.”', prompt='Create a neutral 2:3 product-image placeholder for Christian Science Aurora’s Reading Room shop. Warm white paper texture, thin Aurora teal border, subtle periwinkle corner field, small restrained Morning Gold rule, centered Inter-style typography reading “Product image pending authorization,” with space below for title and format. Minimal, accessible, premium editorial design. No imitation cover art, no publisher logo, no Cross and Crown, no decorative religious symbol.', negative='No fake cover recreation, invented publisher mark, book title unless supplied, decorative cross, halo, 3D mockup distortion, glossy retail effects, or third-party logo.', alt='Product image pending authorized cover use.', release='Written cover-use permission or publisher terms; store the source, rights owner, allowed context, and review date.'),
    dict(group='Reading Room and Events', id='reading-room-exterior', placement='Reading Room Visit — page hero', priority='Launch critical once location confirmed', ratio='16:9', label='Reading Room exterior or entrance', objective='Help visitors recognize the correct public entrance and understand accessibility.', brief='Photograph the actual entrance straight enough to show wayfinding. Include accessible route and hours sign only after verified.', prompt='Using verified reference photographs of the actual Christian Science Reading Room entrance, create an ultra-realistic documentary exterior image. Preserve exact storefront or building entrance, signage, door, ramp or level approach, surrounding context, and materials. Clear natural daylight, uncluttered composition, no people required, 16:9 crop with the entrance immediately understandable.', negative='No invented storefront, altered hours, fake signage, enlarged logo, decorative cross, dramatic sky, visible license plates, artificial glow, or accessibility feature not actually present.', alt='Public entrance to the Christian Science Reading Room.', release='Confirm location, public hours/signage, accessibility route, and trademark permissions.'),
    dict(group='Reading Room and Events', id='event-speaker', placement='Events and public-talk detail pages', priority='Event-specific', ratio='16:9', label='Public talk and audience', objective='Show genuine participation and make event pages feel active without implying endorsements or unverified claims.', brief='Capture the speaker from side/rear audience perspective, plus a second frame of the welcome table or discussion. Preserve event context and avoid close-ups of attendees without releases.', prompt='Using approved event reference photographs, create an ultra-realistic editorial image of a public talk at Christian Science Aurora. A speaker presents naturally at the front while a modest audience is seen from a respectful rear or side angle. Real venue architecture, practical lighting, realistic posture and expressions, no staged applause, calm attentive atmosphere, 16:9 composition with negative space for event title and date.', negative='No fabricated speaker, false attendance size, healing claim text, dramatic spotlight, conference branding, staged applause, identifiable minors, third-party logos, or altered venue.', alt='A public speaker presenting to an audience at a Christian Science event.', release='Speaker agreement, event-media permission, attendee releases where recognizable, and youth restrictions.'),

    dict(group='Articles and About', id='article-arrival', placement='Article — “What can I expect at my first service?”', priority='Launch support', ratio='16:9 or 4:3', label='Visitor entering the church', objective='Support a practical newcomer article with a truthful, low-pressure arrival scene.', brief='Use one adult visitor approaching or entering the actual door, framed from behind or side. The visitor should not look posed or emotionally exaggerated.', prompt='Using actual church entrance references and a released adult model, create an ultra-realistic candid editorial photograph of one visitor approaching the open entrance of First Church of Christ, Scientist, Aurora. Three-quarter rear angle, natural walking motion, exact architecture and signage preserved, soft morning light, understated clothing, calm and practical mood, no direct camera engagement, 16:9 composition with 4:3 crop safety.', negative='No staged smile, dramatic open-door glow, invented entrance, crowd, handshake, praying pose, decorative cross, visible license plate, or artificial sun rays.', alt='A visitor approaching the open entrance of the Aurora church.', release='Adult model and property release; confirm the correct visitor entrance.'),
    dict(group='Articles and About', id='about-exterior-community', placement='About — primary hero', priority='Launch critical', ratio='16:9', label='Church exterior with local context', objective='Present the organization as an established local branch church, distinct from The Mother Church.', brief='Photograph the exact local building with Aurora landscape cues. Avoid visual cues that could imply Boston or a national headquarters.', prompt='Using supplied reference photographs, create an ultra-realistic wide editorial photograph of First Church of Christ, Scientist, Aurora in its real Aurora, Colorado setting. Preserve the exact local building, sign, landscape, neighboring context, and scale. Clear natural daylight, subtle everyday human activity, grounded community feeling, no grand institutional treatment, 16:9 frame with clean space for About-page title.', negative='No Boston architecture, headquarters scale, invented dome or steeple, decorative cross, dramatic clouds, false sign, crowd, or national-church implication.', alt='First Church of Christ, Scientist, Aurora in its local neighborhood setting.', release='Property review, signage verification, and adult releases if recognizable.'),
    dict(group='Articles and About', id='history-archive', placement='About History — archival feature', priority='After historical materials are verified', ratio='4:3', label='Approved historical materials', objective='Preserve local history with reliable source information and an editorial presentation.', brief='Digitize original photos, programs, building records, and clippings at high resolution. Keep original proportions and do not erase meaningful marks. Create captions from verified sources.', prompt='Create a high-resolution editorial flat-lay photograph of verified archival materials from First Church of Christ, Scientist, Aurora: one approved historical building photograph, a dated program, and a simple typed timeline card. Neutral warm-white surface, soft even museum-style light, accurate color, crisp edges, no invented dates or text, no artificial aging, 4:3 composition with each item clearly separated. Use only supplied source scans and preserve them without alteration.', negative='No fabricated historical photo, invented founding date, fake handwriting, artificial sepia, torn-paper effect, decorative cross, or rewritten archival text.', alt='Selected archival materials from the history of First Church of Christ, Scientist, Aurora.', release='Document source, date, owner, caption, copyright status, digitization permission, and approver.'),
    dict(group='Articles and About', id='pastor-books', placement='About — The Christian Science Pastor', priority='Launch support', ratio='4:3', label='The Bible and Science and Health', objective='Explain the Pastor visually with dignity while respecting publication and trademark rules.', brief='Photograph approved editions on a simple reading surface. Keep page text unreadable and avoid arranging books as a decorative religious still life.', prompt='Using approved physical editions and a real church or Reading Room setting, create an ultra-realistic editorial photograph of the Bible and Science and Health with Key to the Scriptures resting naturally together on a simple reading table. Covers shown only as permitted, page text unreadable, soft natural side light, accurate book proportions and colors, warm white background, subtle teal/periwinkle environment accents, 4:3 composition, quiet and informative rather than symbolic.', negative='No glowing books, halo, decorative cross, floating pages, readable copyrighted passages, altered cover design, oversized trademark, flowers or candles added for symbolism, or dramatic spotlight.', alt='The Bible and Science and Health with Key to the Scriptures displayed together.', release='Confirm cover and title-page use, edition, trademark context, and photography permission.'),
]

# ---------- helpers ----------
def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tcMar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tcMar.append(node)
        node.set(qn('w:w'), str(v)); node.set(qn('w:type'), 'dxa')


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = instruction
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)


def add_label_value(doc, label, value, shaded=False):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(1.45)
    table.columns[1].width = Inches(5.85)
    c1, c2 = table.rows[0].cells
    c1.width = Inches(1.45); c2.width = Inches(5.85)
    for c in (c1,c2):
        set_cell_margins(c)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_shading(c1, DEEP_TEAL)
    if shaded: set_cell_shading(c2, MIST)
    p1 = c1.paragraphs[0]; p1.paragraph_format.space_after = Pt(0)
    r = p1.add_run(label.upper()); r.bold = True; r.font.name='Inter'; r.font.size=Pt(8.5); r.font.color.rgb=RGBColor(255,255,255)
    p2 = c2.paragraphs[0]; p2.paragraph_format.space_after = Pt(0)
    r = p2.add_run(value); r.font.name='Inter'; r.font.size=Pt(9.3); r.font.color.rgb=RGBColor.from_string(CHARCOAL)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_prompt_box(doc, heading, text, fill):
    table = doc.add_table(rows=2, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(7.3)
    h, b = table.rows[0].cells[0], table.rows[1].cells[0]
    for c in (h,b): set_cell_margins(c, top=110, start=150, bottom=110, end=150)
    set_cell_shading(h, fill); set_cell_shading(b, WARM)
    p = h.paragraphs[0]; p.paragraph_format.space_after=Pt(0)
    r=p.add_run(heading.upper()); r.bold=True; r.font.name='Inter'; r.font.size=Pt(9); r.font.color.rgb=RGBColor(255,255,255)
    p=b.paragraphs[0]; p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.12
    r=p.add_run(text); r.font.name='Inter'; r.font.size=Pt(9.1); r.font.color.rgb=RGBColor.from_string(CHARCOAL)
    return table


doc=Document()
section=doc.sections[0]
section.top_margin=Inches(.65); section.bottom_margin=Inches(.65); section.left_margin=Inches(.7); section.right_margin=Inches(.7)

styles=doc.styles
normal=styles['Normal']; normal.font.name='Inter'; normal.font.size=Pt(10); normal.font.color.rgb=RGBColor.from_string(CHARCOAL)
normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.15
for name,size,color in [('Title',34,DEEP_TEAL),('Heading 1',24,DEEP_TEAL),('Heading 2',16,DEEP_TEAL),('Heading 3',12,DEEP_PERI)]:
    st=styles[name]; st.font.name='Inter Display'; st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=RGBColor.from_string(color)
    st.paragraph_format.space_before=Pt(8); st.paragraph_format.space_after=Pt(6)

if 'Eyebrow' not in styles:
    st=styles.add_style('Eyebrow',WD_STYLE_TYPE.PARAGRAPH)
else: st=styles['Eyebrow']
st.font.name='Inter'; st.font.size=Pt(8); st.font.bold=True; st.font.color.rgb=RGBColor.from_string(DEEP_PERI)
st.paragraph_format.space_after=Pt(4)

# Header/footer
header=section.header
p=header.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.RIGHT
r=p.add_run('CHRISTIAN SCIENCE AURORA  /  PHOTOGRAPHY SYSTEM'); r.font.name='Inter'; r.font.size=Pt(7.5); r.font.bold=True; r.font.color.rgb=RGBColor.from_string(TEAL)
footer=section.footer
p=footer.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('ChristianScienceAurora.com   •   Phase 1 photography and AI image prompts   •   '); r.font.name='Inter'; r.font.size=Pt(7); r.font.color.rgb=RGBColor.from_string(CHARCOAL)
add_field(p,'PAGE')

# Cover
p=doc.add_paragraph(style='Eyebrow'); p.add_run('CSA  /  CHRISTIAN SCIENCE  /  AURORA')
p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(40); p.paragraph_format.space_after=Pt(8)
r=p.add_run('PHOTOGRAPHY &\nIMAGE PRODUCTION PLAN'); r.font.name='Inter Display'; r.font.size=Pt(33); r.font.bold=True; r.font.color.rgb=RGBColor.from_string(DEEP_TEAL)
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(18)
r=p.add_run('Detailed shot descriptions, production requirements, alt text, permissions safeguards, and AI image prompts for every Phase 1 public-site image placeholder.'); r.font.name='Inter'; r.font.size=Pt(14); r.font.color.rgb=RGBColor.from_string(CHARCOAL)

t=doc.add_table(rows=1, cols=2); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
t.columns[0].width=Inches(3.5); t.columns[1].width=Inches(3.5)
for c,fill in zip(t.rows[0].cells,[TEAL,PERI]): set_cell_shading(c,fill); c.height=Inches(.35)

p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(30)
r=p.add_run('First Church of Christ, Scientist, Aurora, Colorado'); r.bold=True; r.font.name='Inter'; r.font.size=Pt(11)
p=doc.add_paragraph('ChristianScienceAurora.com  |  Phase 1 Checkpoint 2  |  July 2026')
p.runs[0].font.name='Inter'; p.runs[0].font.size=Pt(9); p.runs[0].font.color.rgb=RGBColor.from_string(DEEP_PERI)

doc.add_page_break()

# Executive direction
p=doc.add_paragraph(style='Eyebrow'); p.add_run('01  EXECUTIVE DIRECTION')
doc.add_heading('Use real Aurora photography as the final source of truth.', level=1)
doc.add_paragraph('The public website is designed to feel modern, peaceful, locally grounded, and genuinely welcoming. Final launch photography should document the actual church, Reading Room, people, and community. AI-generated imagery can support storyboarding, crop planning, temporary private mockups, or careful relighting of supplied real photographs, but it should not be presented as documentary evidence of a building, event, congregation, accessibility feature, or historical fact that does not exist.')

add_label_value(doc,'Primary rule','Preserve the approved text-only identity. Do not add a symbol, emblem, decorative cross, halo, glowing hands, or separate Sunday School logo.',True)
add_label_value(doc,'People','Favor natural interaction over direct-to-camera posing. Obtain releases for recognizable adults and strict guardian authorization for youth.',False)
add_label_value(doc,'Rights','Do not show readable Bible Lesson text, unauthorized publication covers, restricted marks, personal information, or unapproved event content.',True)
add_label_value(doc,'Look','Natural Colorado light, realistic contrast, warm white visual field, restrained teal/periwinkle harmony, and minimal filtering.',False)

p=doc.add_paragraph(style='Heading 2'); p.add_run('Launch-priority shot list')
for text in [
    'Tier 1 — Home welcome hero, exterior/entrance, sanctuary wide, Reading Room interior, Sunday School hands, About exterior, and parking/entrance orientation.',
    'Tier 2 — Visitor welcome, Sunday sanctuary detail, Wednesday evening exterior, sermon media, parent/family arrival, article image, and Pastor books.',
    'Tier 3 — Campaign/event photography, archival materials, product-cover system, and new editorial images added as content is approved.',
]:
    p=doc.add_paragraph(style='List Bullet'); p.add_run(text)

p=doc.add_paragraph(style='Heading 2'); p.add_run('Recommended capture package')
for text in [
    'RAW files plus full-resolution edited TIFF or maximum-quality JPEG; retain original metadata privately but strip location metadata from public youth imagery.',
    'For each launch-critical scene, capture horizontal 16:9, standard 4:3, and vertical 4:5 compositions or provide enough resolution for safe crops.',
    'Use consistent filenames: csa_[location]_[subject]_[orientation]_[sequence]_[YYYYMMDD].jpg.',
    'Record asset owner, photographer, date, location, release status, rights restrictions, alt text, crop notes, and next review date in the CMS media record.',
]:
    p=doc.add_paragraph(style='List Bullet'); p.add_run(text)

doc.add_page_break()

# production standards
p=doc.add_paragraph(style='Eyebrow'); p.add_run('02  PRODUCTION STANDARDS')
doc.add_heading('A consistent photographic language', level=1)
standards=[
 ('Lighting','Use real window light, open shade, or balanced practical light. Preserve highlight detail and realistic shadow depth. Avoid fake sunbeams, glowing entrances, and heavy HDR.'),
 ('Composition','Keep architecture straight, allow generous negative space for headings, and protect a centered mobile focal area. Use eye-level visitor perspectives rather than dramatic real-estate angles.'),
 ('People','Capture genuine movement, listening, reading, helping, and conversation. Avoid rows of posed attendees, exaggerated smiles, or imagery implying guaranteed spiritual or physical outcomes.'),
 ('Color','Keep the overall image warm and natural. Teal, periwinkle, and gold may appear as subtle environmental accents; do not force brand colors into clothing or oversaturate the grade.'),
 ('Accessibility','Images must not be the only way practical information is communicated. Parking and entrance visuals need text equivalents and must match verified routes.'),
 ('Youth safeguards','No public child profiles or surnames; no recognizable youth image without documented guardian release; strip geolocation metadata; disable public comments; obtain editorial approval.'),
 ('AI use','For actual locations, always provide reference photographs and instruct the model to preserve exact architecture. Do not publish an AI-invented church, parking route, event, historical scene, or congregation as fact.'),
]
for h,b in standards:
    add_label_value(doc,h,b,shaded=(standards.index((h,b))%2==0))

p=doc.add_paragraph(style='Heading 2'); p.add_run('Global negative prompt')
doc.add_paragraph('No religious clip art, decorative cross, Cross and Crown used as decoration, halo, glowing hands, artificial healing aura, dramatic sky replacement, staged corporate smiles, wellness-spa clichés, exaggerated lens flare, overprocessed HDR, invented architecture, fake event details, readable protected text, unauthorized covers, third-party logos, or identifiable minors without releases.')

# asset sections
current_group=None
asset_num=0
for item in assets:
    if item['group'] != current_group:
        doc.add_page_break()
        current_group=item['group']
        p=doc.add_paragraph(style='Eyebrow'); p.add_run(f'IMAGE GROUP  /  {current_group.upper()}')
        doc.add_heading(current_group, level=1)
        doc.add_paragraph('The following assets are mapped directly to the working Phase 1 public website and its image placeholders.')
    asset_num += 1
    p=doc.add_paragraph(style='Eyebrow'); p.add_run(f'ASSET {asset_num:02d}  /  {item["id"]}')
    doc.add_heading(item['label'], level=2)
    add_label_value(doc,'Placement',item['placement'],True)
    add_label_value(doc,'Priority',item['priority'],False)
    add_label_value(doc,'Crop',item['ratio'],True)
    add_label_value(doc,'Objective',item['objective'],False)
    add_label_value(doc,'Photo brief',item['brief'],True)
    add_prompt_box(doc,'AI generation / transformation prompt',item['prompt'],TEAL)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)
    add_prompt_box(doc,'Negative prompt',item['negative'],DEEP_PERI)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)
    add_label_value(doc,'Alt text',item['alt'],True)
    add_label_value(doc,'Permissions',item['release'],False)
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(10)
    r=p.add_run('DELIVERY: '); r.bold=True; r.font.name='Inter'; r.font.size=Pt(8.5); r.font.color.rgb=RGBColor.from_string(GOLD)
    r=p.add_run(f'{item["id"]}_desktop.webp, {item["id"]}_mobile.webp, full-resolution master, release/rights record, and CMS alt text.'); r.font.name='Inter'; r.font.size=Pt(8.5)

# final checklist
doc.add_page_break()
p=doc.add_paragraph(style='Eyebrow'); p.add_run('FINAL QA')
doc.add_heading('Image approval checklist',level=1)
checks=[
 'The image depicts the actual location, event, person, or approved illustrative concept accurately.',
 'The crop works at desktop, tablet, and mobile sizes without hiding the essential subject.',
 'Alt text describes meaningful content without repeating adjacent copy.',
 'Recognizable adults have releases; recognizable youth have guardian releases and youth approval.',
 'No location metadata remains in public youth files.',
 'No protected text, cover, mark, logo, testimony, or event material appears without permission.',
 'No visual implies guaranteed healing or a medical outcome.',
 'Contrast and overlay areas remain readable at WCAG 2.2 AA targets.',
 'File is exported to AVIF/WebP for the site, with a high-quality source master retained privately.',
 'CMS record includes owner, source, rights, alt text, crops, review date, and publication status.',
]
for c in checks:
    p=doc.add_paragraph(style='List Bullet'); p.add_run('☐ '+c)

p=doc.add_paragraph(style='Heading 2'); p.add_run('Church decisions still needed')
for c in [
 'Preferred visitor entrance and parking route',
 'Confirmed accessible route and accommodations',
 'Reading Room location, hours, phone, entrance, and photography access',
 'Approved adult participants and youth guardian-release process',
 'Permission to photograph visible books, covers, marks, and archival materials',
 'Current campaign/event schedule and speaker-media permissions',
]:
    p=doc.add_paragraph(style='List Bullet'); p.add_run(c)

# document properties
props=doc.core_properties
props.title='Christian Science Aurora Phase 1 Photography and AI Image Prompts'
props.subject='Website photography production plan and image prompts'
props.author='Christian Science Aurora project team'
props.keywords='Christian Science Aurora, website, photography, image prompts, accessibility, permissions'

doc.save(OUT)
print(OUT)
